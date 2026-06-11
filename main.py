"""DocuMind."""

import html
import json
import os
import re
import socket
import smtplib
import threading
from datetime import datetime
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path

import fitz
import gradio as gr
import gspread
from docx import Document
from dotenv import load_dotenv
from flask import Flask, jsonify, request
from groq import Groq
from oauth2client.service_account import ServiceAccountCredentials
from pdf2image import convert_from_path
from PIL import Image
import pytesseract
from rag import build_rag_query
from rag import create_document_index
from rag import format_rag_status
from rag import format_retrieved_evidence
from rag import retrieve_relevant_context


load_dotenv()

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GMAIL_SENDER = os.getenv("GMAIL_SENDER")
GMAIL_APP_PASSWORD = os.getenv("GMAIL_APP_PASSWORD")
GOOGLE_SHEET_ID = os.getenv("GOOGLE_SHEET_ID")

SEMANTIC_ACTION_ITEMS_PROMPT = """You are a business document analyst.
Extract action items semantically from any document type, including meeting notes, emails, reports, resumes, contracts, project plans, and conversational text.
Understand intent and context. Do not rely on exact labels such as Task, Owner, Deadline, or Priority.
Treat natural-language commitments, requests, obligations, follow-ups, review items, approvals, testing work, and deliverables as action items.
Examples include:
- Rachit should prepare the dashboard before Monday
- Follow up with the client next week
- Testing must be completed by Friday
- Send revised proposal to finance team
- John will review the architecture document
Return ONLY valid JSON with this exact shape:
{
  "action_items": [
    {
      "task": "string",
      "owner": "string or null",
      "deadline": "string or null",
      "priority": "High, Medium, Low, or null",
      "confidence": 0.0
    }
  ]
}
Use null for missing owner, deadline, priority, or confidence.
Use an empty action_items array if no action items are present.
Do not include markdown, comments, explanations, or trailing text."""

JSON_REPAIR_PROMPT = """You repair malformed JSON responses.
Return ONLY valid JSON with this exact shape:
{
  "action_items": [
    {
      "task": "string",
      "owner": "string or null",
      "deadline": "string or null",
      "priority": "High, Medium, Low, or null",
      "confidence": 0.0
    }
  ]
}
Use null for missing fields.
Do not add markdown, comments, explanations, or trailing text."""

DOCUMENT_CLASSIFICATION_PROMPT = """You are a document classification analyst.
Classify the document using semantic understanding, not keyword matching.
Choose exactly one document_type from:
meeting_notes, invoice, resume, report, contract, email, research_paper, generic_document.
Return ONLY valid JSON with this exact shape:
{
  "document_type": "meeting_notes",
  "confidence": 0.0,
  "summary": "short summary"
}
Use generic_document if the document type is uncertain.
Do not include markdown, comments, explanations, or trailing text."""

VALID_DOCUMENT_TYPES = {
    "meeting_notes",
    "invoice",
    "resume",
    "report",
    "contract",
    "email",
    "research_paper",
    "generic_document",
}

SPECIALIZED_EXTRACTION_GUIDANCE = {
    "meeting_notes": "Focus on decisions, follow-ups, assigned discussion items, meeting commitments, and due dates.",
    "invoice": "Focus on payment tasks, approvals, billing follow-ups, amount-related actions, and due dates.",
    "resume": "Focus on follow-up actions such as screening, interviewing, reference checks, portfolio review, and hiring decisions.",
    "report": "Focus on recommendations, unresolved issues, follow-ups, review items, and operational next steps.",
    "contract": "Focus on obligations, approvals, signatures, renewals, compliance tasks, and deadlines.",
    "email": "Focus on requests, replies needed, follow-ups, promised work, approvals, and time-sensitive actions.",
    "research_paper": "Focus on experiments, review tasks, citations to check, implementation ideas, and research follow-ups.",
    "generic_document": "Focus on any implied or explicit tasks, responsibilities, follow-ups, approvals, and deadlines.",
}

DOCUMENT_EXTRACTION_SCHEMAS = {
    "meeting_notes": {
        "fields": [
            "task",
            "owner",
            "deadline",
            "priority",
            "confidence",
        ]
    },

    "invoice": {
        "fields": [
            "invoice_number",
            "vendor",
            "invoice_date",
            "due_date",
            "amount",
            "currency",
        ]
    },

    "resume": {
        "fields": [
            "name",
            "email",
            "skills",
            "experience",
            "education",
        ]
    },

    "contract": {
        "fields": [
            "parties",
            "effective_date",
            "expiration_date",
            "obligations",
            "payment_terms",
        ]
    },

    "report": {
        "fields": [
            "title",
            "summary",
            "key_findings",
            "recommendations",
        ]
    },

    "research_paper": {
        "fields": [
            "title",
            "authors",
            "abstract",
            "methods",
            "findings",
        ]
    },

    "email": {
        "fields": [
            "sender",
            "recipient",
            "subject",
            "requests",
            "deadlines",
        ]
    },

    "generic_document": {
        "fields": [
            "summary",
            "important_entities",
            "important_dates",
            "key_points",
        ]
    },
}

GOOGLE_SHEETS_SCOPES = [
    "https://spreadsheets.google.com/feeds",
    "https://www.googleapis.com/auth/drive",
]

flask_app = Flask(__name__)
LAST_EXTRACTION_METHOD = None
LAST_OCR_CONFIDENCE = None
LAST_DOCUMENT_TEXT = ""
LAST_DOCUMENT_ID = None

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}


def generate_workflow_recommendations(
    document_type,
    structured_data=None
):
    workflows = {

        "resume": [
            "Create candidate profile in ATS",
            "Schedule technical interview",
            "Perform reference verification",
            "Generate recruiter summary"
        ],

        "invoice": [
            "Forward to finance team",
            "Create payment reminder",
            "Update accounting records",
            "Flag overdue payments"
        ],

        "contract": [
            "Notify legal stakeholders",
            "Extract key obligations",
            "Schedule renewal reminders",
            "Create CRM record"
        ],

        "report": [
            "Assign review owner",
            "Notify department leads",
            "Track recommendations",
            "Schedule follow-up review"
        ],

        "meeting_notes": [
            "Create task tracker",
            "Assign action items",
            "Schedule follow-up meeting",
            "Notify attendees"
        ],

        "email": [
            "Generate response draft",
            "Create follow-up task",
            "Notify responsible stakeholder",
            "Track pending approvals"
        ],

        "research_paper": [
            "Create research summary",
            "Generate implementation backlog",
            "Assign review owner",
            "Track citations and references"
        ],

        "generic_document": [
            "Generate executive summary",
            "Notify stakeholders",
            "Create follow-up task",
            "Store in document repository"
        ]
    }

    return workflows.get(
        document_type,
        workflows["generic_document"]
    )


# Clean extracted document text into one normalized string.
def clean_extracted_text(text_parts):
    """Normalize extracted text fragments into a single string."""
    return " ".join(" ".join(text.split()) for text in text_parts if text)


# Extract regular text from PDF pages with PyMuPDF.
def extract_pdf_text_with_pymupdf(path):
    """Extract normal page text from a PDF with PyMuPDF."""
    with fitz.open(path) as document:
        return clean_extracted_text([page.get_text() for page in document])


# Extract table content from PDF pages with PyMuPDF table detection.
def extract_pdf_tables_with_pymupdf(path):
    """Extract table cell text from a PDF with PyMuPDF table detection."""
    table_text = []

    with fitz.open(path) as document:
        for page in document:
            table_finder = page.find_tables()
            for table in table_finder.tables:
                for row in table.extract():
                    table_text.append(" ".join(str(cell or "") for cell in row))

    return clean_extracted_text(table_text)


# Extract text and confidence from images with pytesseract OCR.
def extract_text_from_images_with_ocr(images):
    """Extract text and average OCR confidence from a list of images."""
    text_parts = []
    confidence_scores = []

    for index, image in enumerate(images, start=1):
        print(f"OCR processing page/image {index}")
        ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
        page_words = []

        for text, confidence in zip(ocr_data.get("text", []), ocr_data.get("conf", [])):
            text = text.strip()
            if text:
                page_words.append(text)

            try:
                confidence_value = float(confidence)
            except (TypeError, ValueError):
                continue

            if confidence_value >= 0:
                confidence_scores.append(confidence_value)

        text_parts.append(" ".join(page_words))

    average_confidence = round(sum(confidence_scores) / len(confidence_scores), 2) if confidence_scores else 0.0
    return clean_extracted_text(text_parts), average_confidence


# Extract text from PDF pages using OCR.
def extract_pdf_text_with_ocr(path):
    """Extract text from scanned PDF pages with pytesseract OCR."""
    print("OCR activated")
    images = convert_from_path(path)
    print(f"Pages processed: {len(images)}")
    return extract_text_from_images_with_ocr(images)


# Extract plain text from a PDF file path using layered extraction.
def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF using PyMuPDF, table extraction, then OCR."""
    global LAST_EXTRACTION_METHOD, LAST_OCR_CONFIDENCE

    path = Path(pdf_path)

    if not path.is_file():
        return f"Error: PDF file not found at '{pdf_path}'."

    try:
        regular_text = extract_pdf_text_with_pymupdf(path)
        print(f"Extracted {len(regular_text)} characters from PDF with PyMuPDF")

        if len(regular_text) >= 50:
            LAST_EXTRACTION_METHOD = "Text-based PDF"
            LAST_OCR_CONFIDENCE = None
            print("Document type: Text-based PDF")
            return regular_text

        print("Text extraction insufficient. Activating OCR fallback if table extraction also fails.")

        table_text = extract_pdf_tables_with_pymupdf(path)
        combined_text = clean_extracted_text([regular_text, table_text])
        print(f"Extracted {len(combined_text)} characters from PDF with PyMuPDF tables")

        if len(combined_text) >= 50:
            LAST_EXTRACTION_METHOD = "Table-heavy PDF"
            LAST_OCR_CONFIDENCE = None
            print("Document type: Table-heavy PDF")
            return combined_text

        ocr_text, ocr_confidence = extract_pdf_text_with_ocr(path)
        print(f"Extracted {len(ocr_text)} characters from PDF with OCR")
        print(f"OCR confidence: {ocr_confidence}")

        if len(ocr_text) >= 50:
            LAST_EXTRACTION_METHOD = "Scanned/Image PDF"
            LAST_OCR_CONFIDENCE = ocr_confidence
            print("Document type: Scanned/Image PDF")
            return ocr_text

        return "Error: PDF appears to be empty or unreadable after text, table, and OCR extraction."
    except fitz.FileDataError:
        return f"Error: '{pdf_path}' is not a valid PDF file."
    except fitz.EmptyFileError:
        return f"Error: '{pdf_path}' is empty or unreadable."
    except Exception as exc:
        return f"Error: Could not extract text from '{pdf_path}'. {exc}"


# Extract plain text from a TXT file path.
def extract_text_from_txt(txt_path):
    """Extract text from a TXT file and return it as one cleaned string."""
    path = Path(txt_path)

    if not path.is_file():
        return f"Error: Text file not found at '{txt_path}'."

    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="utf-8-sig")
    except Exception as exc:
        return f"Error: Could not extract text from '{txt_path}'. {exc}"

    cleaned_text = " ".join(text.split())
    global LAST_EXTRACTION_METHOD
    LAST_EXTRACTION_METHOD = "TXT"
    print(f"Extracted {len(cleaned_text)} characters from TXT")

    if len(cleaned_text) < 50:
        return "Error: Text file appears to be empty or too short to process."

    return cleaned_text


# Extract plain text from a DOCX file path.
def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file and return it as one cleaned string."""
    path = Path(docx_path)

    if not path.is_file():
        return f"Error: DOCX file not found at '{docx_path}'."

    try:
        document = Document(path)
        paragraph_text = [paragraph.text for paragraph in document.paragraphs]
        cleaned_text = " ".join(" ".join(text.split()) for text in paragraph_text)
        global LAST_EXTRACTION_METHOD
        LAST_EXTRACTION_METHOD = "DOCX"
        print(f"Extracted {len(cleaned_text)} characters from DOCX")

        if len(cleaned_text) < 50:
            return "Error: DOCX file appears to be empty or too short to process."

        return cleaned_text
    except Exception as exc:
        return f"Error: Could not extract text from '{docx_path}'. {exc}"


# Extract plain text from an image file path using OCR.
def extract_text_from_image(image_path):
    """Extract text from an image-based document with pytesseract OCR."""
    global LAST_EXTRACTION_METHOD, LAST_OCR_CONFIDENCE

    path = Path(image_path)

    if not path.is_file():
        return f"Error: Image file not found at '{image_path}'."

    try:
        print("OCR activated")
        with Image.open(path) as image:
            ocr_text, ocr_confidence = extract_text_from_images_with_ocr([image])

        LAST_EXTRACTION_METHOD = "Image OCR"
        LAST_OCR_CONFIDENCE = ocr_confidence
        print("Pages processed: 1")
        print(f"OCR confidence: {ocr_confidence}")
        print(f"Extracted {len(ocr_text)} characters from image with OCR")

        if len(ocr_text) < 50:
            return "Error: Image appears to be empty or unreadable after OCR extraction."

        return ocr_text
    except Exception as exc:
        return f"Error: Could not extract text from image '{image_path}'. {exc}"


# Route supported file types to the correct text extraction function.
def extract_text_from_file(file_path):
    """Extract text from PDF, TXT, DOCX, or image files."""
    suffix = Path(file_path).suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)

    if suffix == ".txt":
        return extract_text_from_txt(file_path)

    if suffix == ".docx":
        return extract_text_from_docx(file_path)

    if suffix in IMAGE_EXTENSIONS:
        return extract_text_from_image(file_path)

    return "Error: Unsupported file type. Please upload a PDF, TXT, DOCX, or image file."


# Normalize LLM output so downstream code always receives valid action-item dictionaries.
def normalize_action_items(parsed_response):
    """Normalize parsed LLM JSON into a list of action-item dictionaries."""
    if isinstance(parsed_response, dict):
        action_items = parsed_response.get("action_items", [])
    elif isinstance(parsed_response, list):
        action_items = parsed_response
    else:
        action_items = []

    normalized_items = []
    for item in action_items:
        if not isinstance(item, dict):
            continue

        normalized_items.append(
            {
                "task": item.get("task") or None,
                "owner": item.get("owner") or None,
                "deadline": item.get("deadline") or None,
                "priority": item.get("priority") or None,
                "confidence": item.get("confidence", item.get("confidence_score")) or None,
            }
        )

    return normalized_items


# Parse JSON from model output and repair it with the LLM if needed.
def parse_or_repair_action_items_json(raw_json, original_text):
    """Parse action-item JSON, asking Groq to repair malformed output if needed."""
    try:
        return json.loads(raw_json)
    except json.JSONDecodeError:
        print("Action item JSON was malformed. Attempting repair with Groq.")

    client = Groq(api_key=os.getenv("GROQ_API_KEY"))
    repair_response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        temperature=0.0,
        max_tokens=1000,
        response_format={"type": "json_object"},
        messages=[
            {
                "role": "system",
                "content": JSON_REPAIR_PROMPT,
            },
            {
                "role": "user",
                "content": (
                    "Original document text:\n"
                    f"{original_text}\n\n"
                    "Malformed extraction response:\n"
                    f"{raw_json}"
                ),
            },
        ],
    )
    repaired_json = repair_response.choices[0].message.content
    return json.loads(repaired_json)


# Normalize the document classification response and fall back when uncertain.
def normalize_classification(parsed_response):
    """Normalize classification JSON into a safe document metadata dictionary."""
    if not isinstance(parsed_response, dict):
        return {
            "document_type": "generic_document",
            "confidence": 0.0,
            "summary": "Unable to classify document.",
        }

    document_type = parsed_response.get("document_type", "generic_document")
    confidence = parsed_response.get("confidence", 0.0)
    summary = parsed_response.get("summary", "")

    try:
        confidence = float(confidence)
    except (TypeError, ValueError):
        confidence = 0.0

    if document_type not in VALID_DOCUMENT_TYPES or confidence < 0.45:
        document_type = "generic_document"

    return {
        "document_type": document_type,
        "confidence": confidence,
        "summary": summary or "No summary available.",
    }


# Classify the extracted document text before action-item extraction.
def classify_document(text):
    """Classify document type with Groq and return type, confidence, and summary."""
    if not GROQ_API_KEY:
        return {"error": "GROQ_API_KEY is missing. Add it to your .env file."}

    if not text or not text.strip():
        return {"error": "No text provided for document classification."}

    raw_json = ""
    try:
        client = Groq(api_key=os.getenv("GROQ_API_KEY"))
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            temperature=0.1,
            max_tokens=500,
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "system",
                    "content": DOCUMENT_CLASSIFICATION_PROMPT,
                },
                {
                    "role": "user",
                    "content": text,
                },
            ],
        )

        raw_json = response.choices[0].message.content
        classification = normalize_classification(json.loads(raw_json))
        print(f"Classification result: {classification}")
        print(f"Confidence score: {classification['confidence']}")
        return classification
    except json.JSONDecodeError as exc:
        print(f"Classification JSON parsing failed: {exc}")
        classification = {
            "document_type": "generic_document",
            "confidence": 0.0,
            "summary": "Classification failed, defaulted to generic_document.",
        }
        print(f"Classification result: {classification}")
        print(f"Confidence score: {classification['confidence']}")
        return classification
    except Exception as exc:
        print(f"Classification failed: {exc}")
        classification = {
            "document_type": "generic_document",
            "confidence": 0.0,
            "summary": "Classification failed, defaulted to generic_document.",
        }
        print(f"Classification result: {classification}")
        print(f"Confidence score: {classification['confidence']}")
        return classification
    
def summarize_document(text):
    """Generate a semantic document summary."""

    prompt = f"""
    You are an intelligent document summarization system.

    Analyze the document and provide:
    - concise summary
    - key insights
    - important entities
    - important dates
    - major conclusions

    Return ONLY valid JSON.

    Document:
    {text[:12000]}
    """

    try:

        client = Groq(
            api_key=os.getenv("GROQ_API_KEY")
        )

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            temperature=0.2,
            max_tokens=1200,
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "system",
                    "content": "You summarize documents intelligently."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
        )

        return json.loads(
            response.choices[0].message.content
        )

    except Exception as exc:

        return {
            "error": "Document summarization failed.",
            "details": str(exc),
        }

def answer_document_question(question):
    """Answer questions using the uploaded document."""

    global LAST_DOCUMENT_TEXT

    if not LAST_DOCUMENT_TEXT.strip():

        return (
            "No document has been uploaded yet."
        )

    if not question.strip():

        return (
            "Please enter a question."
        )

    prompt = f"""
    You are an intelligent document assistant.

    Answer the user's question ONLY
    using the provided document.

    If the answer is not present
    in the document, say:
    "The document does not contain that information."

    Document:
    {LAST_DOCUMENT_TEXT[:12000]}

    Question:
    {question}
    """

    try:

        client = Groq(
            api_key=os.getenv("GROQ_API_KEY")
        )

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            temperature=0.2,
            max_tokens=1000,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You answer questions "
                        "about documents."
                    ),
                },
                {
                    "role": "user",
                    "content": prompt,
                },
            ],
        )

        return response.choices[
            0
        ].message.content

    except Exception as exc:

        return (
            f"Question answering failed: "
            f"{str(exc)}"
        )


def answer_document_question(question):
    """Answer questions using the uploaded document."""

    global LAST_DOCUMENT_TEXT

    if not LAST_DOCUMENT_TEXT.strip():

        return (
            "No document has been uploaded yet."
        )

    if not question.strip():

        return (
            "Please enter a question."
        )

    prompt = f"""
    You are an intelligent document assistant.

    Answer the user's question ONLY
    using the provided document.

    If the answer is not present
    in the document, say:
    "The document does not contain that information."

    Document:
    {LAST_DOCUMENT_TEXT[:12000]}

    Question:
    {question}
    """

    try:

        client = Groq(
            api_key=os.getenv("GROQ_API_KEY")
        )

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            temperature=0.2,
            max_tokens=1000,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You answer questions "
                        "about documents."
                    ),
                },
                {
                    "role": "user",
                    "content": prompt,
                },
            ],
        )

        return response.choices[
            0
        ].message.content

    except Exception as exc:

        return (
            f"Question answering failed: "
            f"{str(exc)}"
        )

# Build a classification-aware semantic extraction prompt.
def build_action_items_prompt(classification):
    """Build a specialized semantic extraction prompt from document classification."""
    document_type = classification.get("document_type", "generic_document")
    guidance = SPECIALIZED_EXTRACTION_GUIDANCE.get(document_type, SPECIALIZED_EXTRACTION_GUIDANCE["generic_document"])
    summary = classification.get("summary", "")

    return (
        f"{SEMANTIC_ACTION_ITEMS_PROMPT}\n\n"
        f"Detected document_type: {document_type}\n"
        f"Document summary: {summary}\n"
        f"Specialized extraction guidance: {guidance}"
    )

def build_dynamic_extraction_prompt(document_type):
    """Build a dynamic extraction prompt based on document type."""

    schema = DOCUMENT_EXTRACTION_SCHEMAS.get(
        document_type,
        DOCUMENT_EXTRACTION_SCHEMAS["generic_document"]
    )

    fields = schema["fields"]

    json_template = {
        field: "string, list, or null"
        for field in fields
    }

    return f"""
    You are an intelligent document extraction system.

    Extract structured information from the document semantically.

    Document type:
    {document_type}

    Return ONLY valid JSON.

    Required JSON structure:
    {json.dumps(json_template, indent=2)}

    Use semantic understanding.
    Do not rely on fixed keywords.
    Use null when information is unavailable.
    """


# Use Groq to convert document text into structured action items.
def extract_action_items(text, classification=None):
    """Extract structured action items from document text using Groq."""
    if not GROQ_API_KEY:
        return {"error": "GROQ_API_KEY is missing. Add it to your .env file."}

    if not text or not text.strip():
        return {"error": "No text provided for action item extraction."}

    classification = classification or {
        "document_type": "generic_document",
        "confidence": 0.0,
        "summary": "No classification provided.",
    }
    system_prompt = build_action_items_prompt(classification)
    raw_json = ""

    try:
        client = Groq(api_key=os.getenv("GROQ_API_KEY"))
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            temperature=0.1,
            max_tokens=1000,
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "system",
                    "content": system_prompt,
                },
                {
                    "role": "user",
                    "content": text,
                }
            ],
        )

        raw_json = response.choices[0].message.content
        parsed_response = parse_or_repair_action_items_json(raw_json, text)
        return normalize_action_items(parsed_response)
    except json.JSONDecodeError as exc:
        print(f"Action item JSON repair failed: {exc}")
        return []
    except Exception as exc:
        return {
            "error": "Could not extract action items.",
            "details": str(exc),
        }

def extract_structured_data(text, classification):
    """Extract dynamic structured data based on document type."""

    document_type = classification.get(
        "document_type",
        "generic_document"
    )

    prompt = build_dynamic_extraction_prompt(document_type)

    try:
        client = Groq(api_key=os.getenv("GROQ_API_KEY"))

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            temperature=0.1,
            max_tokens=1500,
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "system",
                    "content": prompt,
                },
                {
                    "role": "user",
                    "content": text,
                },
            ],
        )

        structured_data = json.loads(
            response.choices[0].message.content
        )

        return structured_data

    except Exception as exc:
        return {
            "error": "Structured extraction failed.",
            "details": str(exc),
        }

# Append extracted action items to the configured Google Sheet.
def log_to_google_sheet(action_items):
    """Append action items to the configured Google Sheet."""
    credentials_path = Path("credentials.json")

    if not GOOGLE_SHEET_ID:
        return {"error": "GOOGLE_SHEET_ID is missing. Add it to your .env file."}

    if not credentials_path.is_file():
        return {"error": "credentials.json was not found."}

    if isinstance(action_items, dict) and "error" in action_items:
        return {"error": "Action items contain an error and were not logged.", "details": action_items}

    if not isinstance(action_items, list):
        return {"error": "Action items must be a list of dictionaries."}

    try:
        credentials = ServiceAccountCredentials.from_json_keyfile_name(
            str(credentials_path),
            GOOGLE_SHEETS_SCOPES,
        )
        client = gspread.authorize(credentials)
        worksheet = client.open_by_key(GOOGLE_SHEET_ID).sheet1

        logged_at = datetime.now().isoformat(timespec="seconds")
        rows = [
            [
                item.get("task", ""),
                item.get("owner", ""),
                item.get("deadline", ""),
                item.get("priority", ""),
                logged_at,
            ]
            for item in action_items
            if isinstance(item, dict)
        ]

        if not rows:
            return {"error": "No valid action items found to log."}

        worksheet.append_rows(rows, value_input_option="USER_ENTERED")
        return {"success": True, "rows_logged": len(rows)}
    except gspread.exceptions.SpreadsheetNotFound:
        return {"error": "Google Sheet was not found. Check GOOGLE_SHEET_ID and sharing permissions."}
    except gspread.exceptions.APIError as exc:
        return {"error": "Google Sheets API request failed.", "details": str(exc)}
    except Exception as exc:
        return {"error": "Google Sheets authentication or logging failed.", "details": str(exc)}


# Send extracted action items as an HTML email summary.
# Send structured document summaries as dynamic HTML emails.
def send_summary_email(document_data, classification, recipient_email):
    """Send dynamic document summaries via email."""

    if not GMAIL_SENDER:
        return {
            "error": "GMAIL_SENDER is missing."
        }

    if not GMAIL_APP_PASSWORD:
        return {
            "error": "GMAIL_APP_PASSWORD is missing."
        }

    if not recipient_email:
        return {
            "error": "Recipient email is required."
        }

    document_type = classification.get(
        "document_type",
        "generic_document"
    )

    subject = (
        f"Document Intelligence Summary - "
        f"{document_type.replace('_', ' ').title()}"
    )

    # Build HTML dynamically
    html_rows = ""

    if isinstance(document_data, list):

        for item in document_data:

            if not isinstance(item, dict):
                continue

            for key, value in item.items():

                html_rows += f"""
                <tr>
                    <td><b>{html.escape(str(key))}</b></td>
                    <td>{html.escape(str(value))}</td>
                </tr>
                """

    elif isinstance(document_data, dict):

        for key, value in document_data.items():

            html_rows += f"""
            <tr>
                <td><b>{html.escape(str(key))}</b></td>
                <td>{html.escape(str(value))}</td>
            </tr>
            """

    html_body = f"""
    <html>
        <body>

            <h2>
                Document Intelligence Summary
            </h2>

            <p>
                <b>Document Type:</b>
                {document_type}
            </p>

            <table
                border="1"
                cellpadding="8"
                cellspacing="0"
                style="border-collapse: collapse;"
            >
                <tbody>
                    {html_rows}
                </tbody>
            </table>

        </body>
    </html>
    """

    message = MIMEMultipart("alternative")

    message["From"] = GMAIL_SENDER
    message["To"] = recipient_email
    message["Subject"] = subject

    message.attach(
        MIMEText(html_body, "html")
    )

    try:

        with smtplib.SMTP(
            "smtp.gmail.com",
            587
        ) as server:

            server.starttls()

            server.login(
                GMAIL_SENDER,
                GMAIL_APP_PASSWORD
            )

            server.sendmail(
                GMAIL_SENDER,
                recipient_email,
                message.as_string()
            )

        return {
            "success": True,
            "email_sent_to": recipient_email,
        }

    except Exception as exc:

        return {
            "error": "Email sending failed.",
            "details": str(exc),
        }


# Convert action item dictionaries into rows Gradio can display.
def format_action_items_for_table(action_items):
    """Convert action item dictionaries into table rows for the UI."""
    return [
        [
            item.get("task", ""),
            item.get("owner", ""),
            item.get("deadline", ""),
            item.get("priority", ""),
            item.get("confidence", ""),
        ]
        for item in action_items
        if isinstance(item, dict)
    ]


# Summarize action item counts by priority for the Gradio dashboard.
def format_priority_dashboard(action_items):
    """Return a Markdown summary of action item counts by priority."""
    counts = {
        "High": 0,
        "Medium": 0,
        "Low": 0,
    }

    for item in action_items:
        if not isinstance(item, dict):
            continue

        priority = str(item.get("priority", "")).strip().title()
        if priority in counts:
            counts[priority] += 1

    high_label = "task" if counts["High"] == 1 else "tasks"
    medium_label = "task" if counts["Medium"] == 1 else "tasks"
    low_label = "task" if counts["Low"] == 1 else "tasks"

    return (
        f"### Priority Dashboard\n\n"
        f"🔴 High Priority: {counts['High']} {high_label}\n\n"
        f"🟡 Medium Priority: {counts['Medium']} {medium_label}\n\n"
        f"🟢 Low Priority: {counts['Low']} {low_label}"
    )


# Format document classification metadata for the Gradio UI.
def format_classification_summary(classification):
    """Return a Markdown summary of document classification metadata."""
    document_type = classification.get("document_type", "generic_document")
    confidence = classification.get("confidence", 0.0)
    summary = classification.get("summary", "No summary available.")

    return (
        f"### Document Classification\n\n"
        f"**Type:** `{document_type}`\n\n"
        f"**Confidence:** {confidence:.2f}\n\n"
        f"**Summary:** {summary}"
    )


def run_document_pipeline(pdf_path, recipient_email):
    """Run the full document processing workflow and return structured results."""
    global LAST_DOCUMENT_ID, LAST_DOCUMENT_TEXT

    page_count = None

    if pdf_path.lower().endswith(".pdf"):
        try:
            doc = fitz.open(pdf_path)
            page_count = len(doc)
            doc.close()
        except Exception:
            page_count = None

    extracted_text = extract_text_from_file(pdf_path)

    if extracted_text.startswith("Error:"):
        return {"status": "error", "message": extracted_text}

    LAST_DOCUMENT_ID = f"{Path(pdf_path).stem}-{abs(hash(pdf_path))}"
    LAST_DOCUMENT_TEXT = extracted_text

    print("Extracted raw text:")
    print(extracted_text)

    rag_index_result = create_document_index(
        LAST_DOCUMENT_ID,
        extracted_text
    )
    print(f"RAG index result: {rag_index_result}")

    classification = classify_document(extracted_text)

    if isinstance(classification, dict) and "error" in classification:
        return {"status": "error", "message": classification["error"]}

    rag_query = build_rag_query(
        classification.get("document_type", "generic_document")
    )
    rag_result = retrieve_relevant_context(
        LAST_DOCUMENT_ID,
        rag_query
    )
    print(f"RAG retrieval query: {rag_query}")
    print(f"RAG retrieved chunks: {len(rag_result.get('chunks', []))}")

    llm_context = (
        rag_result.get("context")
        if rag_result.get("context")
        else extracted_text
    )

    document_summary = summarize_document(
        llm_context
    )

    # Initialize variables safely
    structured_data = None
    action_items = []

    document_type = classification.get(
        "document_type",
        "generic_document"
    )
    workflow_recommendations = (
        generate_workflow_recommendations(
            document_type,
            structured_data
        )
)

    # Use action item extraction only for meeting notes
    if document_type == "meeting_notes":

        action_items = extract_action_items(
            llm_context,
            classification
        )

        if isinstance(action_items, dict) and "error" in action_items:
            return {
                "status": "error",
                "message": action_items["error"]
            }

    # Use dynamic structured extraction for all other document types
    else:

        structured_data = extract_structured_data(
            llm_context,
            classification
        )

        print("Structured Data:")
        print(json.dumps(structured_data, indent=2))

        if isinstance(structured_data, dict) and "error" in structured_data:
            return {
                "status": "error",
                "message": structured_data["error"]
            }

    # No action items found
    if not action_items and not structured_data:

        return {
            "status": "success",
            "page_count": page_count,
            "workflow_recommendations":
                workflow_recommendations,
            "action_items": [],
            "structured_data": structured_data,
            "classification": classification,
            "document_summary": document_summary,
            "rag_index": rag_index_result,
            "rag_result": rag_result,
            "extraction_method": LAST_EXTRACTION_METHOD,
            "ocr_confidence": LAST_OCR_CONFIDENCE,
            "sheet_result": {
                "success": True,
                "rows_logged": 0
            },
            "email_result": {
                "success": True,
                "email_sent_to": recipient_email,
                "skipped": True
            },
        }

    # Only log/email meeting note action items
    if action_items:

        sheet_result = log_to_google_sheet(action_items)

        if not sheet_result.get("success"):
            return {
                "status": "error",
                "message": sheet_result.get(
                    "error",
                    "Google Sheets logging failed."
                ),
            }

        email_payload = (
        action_items
        if action_items
        else structured_data
        )

        email_result = send_summary_email(
        email_payload,
        classification,
        recipient_email
        )

        if not email_result.get("success"):
            return {
                "status": "error",
                "message": email_result.get(
                    "error",
                    "Email sending failed."
                ),
            }

    else:

        sheet_result = {
            "success": True,
            "rows_logged": 0,
            "skipped": True,
        }

        email_result = {
            "success": True,
            "email_sent_to": recipient_email,
            "skipped": True,
        }

        print(
            "Workflow Recommendations:"
        )

        print(
            workflow_recommendations
      )

    return {
        "status": "success",
        "page_count": page_count,
        "workflow_recommendations":
            workflow_recommendations,
        "action_items": action_items,
        "structured_data": structured_data,
        "classification": classification,
        "document_summary": document_summary,
        "rag_index": rag_index_result,
        "rag_result": rag_result,
        "extraction_method": LAST_EXTRACTION_METHOD,
        "ocr_confidence": LAST_OCR_CONFIDENCE,
        "sheet_result": sheet_result,
        "email_result": email_result,
    }

# Handle Gradio submissions and shape output for the UI.
# Handle Gradio submissions and shape output for the UI.
def process_document(pdf_file, recipient_email):
    """Process a document, log action items, email the summary, and return UI output."""

    if not pdf_file:
        return "", "", "", "", "", [], (
            "Error: Please upload a PDF, TXT, DOCX, or image file."
        )

    result = run_document_pipeline(
        pdf_file,
        recipient_email
    )

    if result["status"] == "error":
        return "", "", "", "", "", [], (
            f"Error: {result['message']}"
        )

    classification_summary = format_classification_summary(
        result["classification"]
    )

    # Document summary
    document_summary = result.get(
        "document_summary"
    )

    if document_summary is None:
        summary_output = "{}"
    else:
        summary_output = json.dumps(
            document_summary,
            indent=2
        )

    # Structured extraction output
    structured_data = result.get(
        "structured_data"
    )

    if structured_data is None:
        structured_output = "{}"
    else:
        structured_output = json.dumps(
            structured_data,
            indent=2
        )

    evidence_output = format_retrieved_evidence(
        result.get("rag_result", {})
    )
    rag_status_output = format_rag_status(
        result.get("rag_index", {}),
        result.get("rag_result", {})
    )

    # Show table only for meeting notes
    document_type = result["classification"][
        "document_type"
    ]

    if document_type == "meeting_notes":

        table_rows = format_action_items_for_table(
            result["action_items"]
        )

    else:

        table_rows = []

    # Status handling
    if result["email_result"].get("skipped"):

        status = (
            f"Success: Extraction method: "
            f"{result['extraction_method']}. "
            f"OCR confidence: "
            f"{result['ocr_confidence']}. "
            f"Document type: "
            f"{result['classification']['document_type']}. "
            "No action items found, so "
            "Google Sheets logging and email "
            "sending were skipped."
        )

        return (
            classification_summary,
            rag_status_output,
            summary_output,
            structured_output,
            evidence_output,
            table_rows,
            status,
        )

    status = (
        f"Success: Extraction method: "
        f"{result['extraction_method']}. "
        f"OCR confidence: "
        f"{result['ocr_confidence']}. "
        f"Document type: "
        f"{result['classification']['document_type']}. "
        f"Logged "
        f"{result['sheet_result']['rows_logged']} "
        f"rows to Google Sheets. "
        f"Email sent to "
        f"{result['email_result']['email_sent_to']}."
    )

    return (
        classification_summary,
        rag_status_output,
        summary_output,
        structured_output,
        evidence_output,
        table_rows,
        status,
    )


# Report whether the Flask API is running.
@flask_app.get("/")
def api_home():
    """Return a simple API status response for browser checks."""
    return jsonify(
        {
            "status": "running",
            "health": "/health",
            "webhook": "/webhook",
            "gradio_ui": "http://127.0.0.1:7860",
        }
    )


# Report whether the Flask API is running.
@flask_app.get("/health")
def health():
    """Return a basic health check response."""
    return jsonify({"status": "running"})


# Accept webhook requests and process a PDF document.
@flask_app.post("/webhook")
def webhook():
    """Process documents from a JSON webhook request."""
    payload = request.get_json(silent=True) or {}
    pdf_path = payload.get("pdf_path")
    recipient_email = payload.get("recipient_email")

    if not pdf_path:
        return jsonify({"status": "error", "message": "pdf_path is required."}), 400

    if not recipient_email:
        return jsonify({"status": "error", "message": "recipient_email is required."}), 400

    result = run_document_pipeline(pdf_path, recipient_email)

    if result["status"] == "error":
        return jsonify({"status": "error", "message": result["message"]}), 500

    return jsonify(
        {
            "status": "success",
            "classification": result["classification"],
            "ocr_confidence": result["ocr_confidence"],
            "rag_index": result.get("rag_index"),
            "rag_chunks": result.get("rag_result", {}).get("chunks", []),
            "action_items": result["action_items"],
            "structured_data": result.get("structured_data"),
        }
    )


# Start the Flask app on port 5000.
def run_flask_app():
    """Run Flask without blocking the Gradio UI."""
    flask_app.run(host="127.0.0.1", port=5000, debug=False, use_reloader=False)


# Start Flask in a separate thread so Gradio can run at the same time.
def start_flask_thread():
    """Start the Flask API in a daemon thread."""
    flask_thread = threading.Thread(target=run_flask_app, daemon=True)
    flask_thread.start()
    return flask_thread


# Answer simple word-count questions from the full extracted text.
def answer_count_question(question):
    """Answer word-count questions deterministically instead of using RAG."""
    patterns = [
        r"(?:how many|number of|no of|count of|times of)\s+['\"]?([a-zA-Z0-9_-]+)['\"]?\s+(?:is\s+)?(?:used|appears|mentioned)",
        r"['\"]?([a-zA-Z0-9_-]+)['\"]?\s+(?:is\s+)?(?:used|appears|mentioned)\s+(?:how many|number of|no of|count of)?",
    ]

    lowered_question = question.lower()
    for pattern in patterns:
        match = re.search(pattern, lowered_question)
        if not match:
            continue

        word = match.group(1)
        count = len(re.findall(rf"\b{re.escape(word)}\b", LAST_DOCUMENT_TEXT, flags=re.IGNORECASE))
        return f'The word "{word}" is used {count} times in this document.'

    return None


# Answer user questions with retrieved document chunks instead of the whole document.
def answer_document_question(question):
    """Answer questions using RAG over the uploaded document."""
    if not LAST_DOCUMENT_TEXT.strip() or not LAST_DOCUMENT_ID:
        return "No document has been uploaded yet."

    if not question.strip():
        return "Please enter a question."

    count_answer = answer_count_question(question)
    if count_answer:
        return count_answer

    retrieval_result = retrieve_relevant_context(
        LAST_DOCUMENT_ID,
        question,
        top_k=5
    )
    context = (
        retrieval_result.get("context")
        if retrieval_result.get("context")
        else LAST_DOCUMENT_TEXT[:12000]
    )

    prompt = f"""
    You are an intelligent document assistant.

    Answer the user's question ONLY using the retrieved document context.
    If the answer is not present in the context, say:
    "The document does not contain that information."

    Retrieved context:
    {context}

    Question:
    {question}
    """
    try:
        client = Groq(api_key=os.getenv("GROQ_API_KEY"))

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            temperature=0.2,
            max_tokens=1000,
            messages=[
                {
                    "role": "system",
                    "content": "You answer questions about documents using retrieved evidence.",
                },
                {
                    "role": "user",
                    "content": prompt,
                },
            ],
        )

        answer = response.choices[0].message.content

        sources = [
            chunk.get(
                "source",
                f"Chunk {chunk.get('chunk_index', '?')}"
            )
            for chunk in retrieval_result.get(
                "chunks",
                []
            )
        ]

        unique_sources = list(
            dict.fromkeys(sources)
        )

        return (
            answer
            + "\n\n---\n\n"
            + "**Retrieved Evidence**\n\n"
            + "\n".join(
                f"- {source}"
                for source in unique_sources
            )
        )

    except Exception as exc:
        return f"Question answering failed: {str(exc)}"


# Build the Gradio web interface.
def build_ui():
    print("USING BUILD_UI #1")
    """Build the Gradio interface."""

    with gr.Blocks(
        title="AI Document Intelligence Platform",
        theme=gr.themes.Soft()
    ) as demo:

        gr.Markdown("""
# 📄 DocuMind

### AI-Powered Document Intelligence & Workflow Automation

Turn unstructured documents into searchable knowledge, actionable insights, and business workflows.

---

### Capabilities

🔍 OCR & Document Processing

🧠 AI Classification & Summarization

📊 Structured Data Extraction

✅ Action Item Detection

💬 Chat With Your Documents

📚 Semantic Search & RAG

⚙️ Workflow Recommendations

---

Supported Formats:
PDF • DOCX • TXT • PNG • JPG

Powered by:
Llama 3.3 70B • ChromaDB • Sentence Transformers

---
    
Upload a document below to begin analysis.
                    
""")

        with gr.Row():

            pdf_input = gr.File(
                label="📄 Upload Document",
                file_types=[
                    ".pdf",
                    ".txt",
                    ".docx",
                    ".png",
                    ".jpg",
                    ".jpeg",
                    ".tif",
                    ".tiff",
                    ".bmp",
                    ".webp",
                ],
                type="filepath",
            )

            recipient_input = gr.Textbox(
                label="📧 Recipient Email Address",
                placeholder="recipient@example.com"
            )

        question_input = gr.Textbox(
            label="💬 Ask Questions About Document",
            placeholder="Example: What are the key findings?"
        )

        with gr.Row():

            process_button = gr.Button(
                "🚀 Analyze Document",
                variant="primary"
            )

            question_button = gr.Button(
                "💬 Chat with Document",
                variant="secondary"
            )

        with gr.Tabs():

            # -------------------------------
            # Classification Tab
            # -------------------------------
            with gr.Tab("📄 Classification"):

                with gr.Accordion(
                    "View Document Classification",
                    open=True
                ):

                    classification_output = gr.Markdown(
                        label="Document Classification"
                    )

            # -------------------------------
            # Summary Tab
            # -------------------------------
            with gr.Tab("📝 Summary"):

                with gr.Accordion(
                    "View AI Summary",
                    open=True
                ):

                    summary_output = gr.Code(
                        label="Document Summary",
                        language="json"
                    )

            # -------------------------------
            # Structured Data Tab
            # -------------------------------
            with gr.Tab("🧠 Structured Data"):

                with gr.Accordion(
                    "View Extracted Data",
                    open=True
                ):

                    structured_data_output = gr.Code(
                        label="Structured Document Data",
                        language="json"
                    )

            # -------------------------------
            # RAG Evidence Tab
            # -------------------------------
            with gr.Tab("RAG Evidence"):

                with gr.Accordion(
                    "View Retrieved Context",
                    open=True
                ):

                    rag_evidence_output = gr.Markdown(
                        label="Retrieved Evidence"
                    )

            # -------------------------------
            # Action Items Tab
            # -------------------------------
            with gr.Tab("✅ Action Items"):

                with gr.Accordion(
                    "View Action Items",
                    open=True
                ):

                    priority_dashboard_output = gr.Markdown(
                        label="Priority Dashboard"
                    )

                    action_items_output = gr.Dataframe(
                        headers=[
                            "Task",
                            "Owner",
                            "Deadline",
                            "Priority",
                            "Confidence"
                        ],
                        label="Extracted Action Items",
                        interactive=False,
                    )

            # -------------------------------
            # AI Chat Tab
            # -------------------------------
            with gr.Tab("💬 AI Chat"):

                with gr.Accordion(
                    "Chat With Your Document",
                    open=True
                ):

                    answer_output = gr.Markdown(
                        label="AI Answer"
                    )

            # -------------------------------
            # Status Tab
            # -------------------------------
            with gr.Tab("📌 Status"):

                with gr.Accordion(
                    "Processing Status",
                    open=True
                ):

                    status_output = gr.Textbox(
                        label="Status",
                        interactive=False
                    )

        # -------------------------------
        # Process Button Logic
        # -------------------------------
        process_button.click(
            fn=process_document,
            inputs=[
                pdf_input,
                recipient_input
            ],
            outputs=[
                classification_output,
                summary_output,
                structured_data_output,
                rag_evidence_output,
                priority_dashboard_output,
                action_items_output,
                status_output,
            ],
        )

        # -------------------------------
        # Question Answering Logic
        # -------------------------------
        question_button.click(
            fn=answer_document_question,
            inputs=[question_input],
            outputs=[answer_output],
        )

    return demo


# Build a simplified interview-friendly Gradio interface.
def build_ui():
    print("USING BUILD_UI #2")
    """Build a simple Gradio interface with only essential sections."""
    with gr.Blocks(title="DocuMind") as demo:
        gr.Markdown("""
# 📄 DocuMind

### AI-Powered Document Intelligence & Workflow Automation

Turn unstructured documents into searchable knowledge, actionable insights, and business workflows.

---

### Capabilities

🔍 OCR & Document Processing

🧠 AI Classification & Summarization

📊 Structured Data Extraction

✅ Action Item Detection

💬 Chat With Your Documents

📚 Semantic Search & RAG

⚙️ Workflow Recommendations

---

Supported Formats:
PDF • DOCX • TXT • PNG • JPG

Powered by:
Llama 3.3 70B • ChromaDB • Sentence Transformers

---
    
Upload a document below to begin analysis.
                    
""")
        
        with gr.Row():

              gr.Markdown("""
              ### 📂 Supported Formats
              PDF, DOCX, TXT, Images
              """)

        gr.Markdown("""
        ### 🤖 AI Engine
        Llama 3.3 70B
        """)

        gr.Markdown("""
        ### 🔍 Search
        ChromaDB + RAG
        """)

        with gr.Row():
            pdf_input = gr.File(
                label="Upload Document",
                file_types=[
                    ".pdf",
                    ".txt",
                    ".docx",
                    ".png",
                    ".jpg",
                    ".jpeg",
                    ".tif",
                    ".tiff",
                    ".bmp",
                    ".webp",
                ],
                type="filepath",
            )

            with gr.Accordion(
    "Advanced Options",
    open=False
        ): 
               recipient_input = gr.Textbox(
                label="Recipient Email Address",
                placeholder="recipient@example.com",
            )

        process_button = gr.Button("Analyze Document", variant="primary")

        with gr.Tabs():
            with gr.Tab("Results"):
                with gr.Row():
                    classification_output = gr.Markdown(label="Document Classification")
                    rag_status_output = gr.Markdown(label="RAG Status")

                status_output = gr.Textbox(label="Processing Status", interactive=False)

                with gr.Accordion("Summary", open=True):
                    summary_output = gr.Code(label="Document Summary", language="json")

                with gr.Accordion("Structured Data", open=True):
                    structured_data_output = gr.Code(label="Extracted Data", language="json")

                with gr.Accordion("Action Items", open=False):
                    action_items_output = gr.Dataframe(
                        headers=["Task", "Owner", "Deadline", "Priority", "Confidence"],
                        label="Action Items",
                        interactive=False,
                    )

                with gr.Accordion("Retrieved Context Used By RAG", open=False):
                    rag_evidence_output = gr.Markdown(label="Retrieved Context")

            with gr.Tab("Chat with Document"):
                question_input = gr.Textbox(
                    label="Ask a question about the uploaded document",
                    placeholder="Example: What are the key findings?",
                )
                question_button = gr.Button("Chat with Document", variant="secondary")
                answer_output = gr.Markdown(label="AI Answer")

        process_button.click(
            fn=process_document,
            inputs=[pdf_input, recipient_input],
            outputs=[
                classification_output,
                rag_status_output,
                summary_output,
                structured_data_output,
                rag_evidence_output,
                action_items_output,
                status_output,
            ],
        )

        question_button.click(
            fn=answer_document_question,
            inputs=[question_input],
            outputs=[answer_output],
        )

    return demo

if __name__ == "__main__":

    start_flask_thread()

    gradio_app = build_ui()

    def find_free_port(start=7860, end=7900):

        for port in range(start, end):

            with socket.socket(
                socket.AF_INET,
                socket.SOCK_STREAM
            ) as s:

                try:

                    s.bind(("127.0.0.1", port))

                    return port

                except OSError:

                    continue

        raise OSError(
            "No free port found in range 7860-7900"
        )

    free_port = find_free_port()

    print(
        f"Starting Gradio on port {free_port}"
    )

    gradio_app.launch(
        server_name="127.0.0.1",
        server_port=free_port,
        share=False
    )
